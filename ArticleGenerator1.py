import traceback
import requests
import streamlit as st
import os
import json
import warnings
import io
from datetime import datetime, timedelta
from crewai import Agent, Task, Crew
from docx import Document
from docx.shared import Pt
from docx import Document as DocxDocument  # Import for reading Word documents
from langchain.chat_models import AzureChatOpenAI
import openai

# Suppress warnings
warnings.filterwarnings('ignore')

# Helper function to load and save configurations
def load_config():
    try:
        with open("agent_task_config.json", "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}

def save_config(config):
    with open("agent_task_config.json", "w") as f:
        json.dump(config, f)

# Function to read content from a Word document
def read_docx(file):
    doc = DocxDocument(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Load persisted configurations at startup
config = load_config()

# Streamlit UI
st.title("Research Article Generator")

# Azure OpenAI API inputs
azure_api_key = st.text_input("Enter your Azure OpenAI API Key", type="password")
azure_api_base = "https://rstapestryopenai2.openai.azure.com/"
azure_api_version = "2024-02-15-preview"
deployment_name = "gpt-4"

# Initialize Azure OpenAI instance
if azure_api_key:
    try:
        openai.api_key = azure_api_key
        openai.api_base = azure_api_base

        llm = AzureChatOpenAI(
            openai_api_key=azure_api_key,
            openai_api_base=azure_api_base,
            openai_api_version=azure_api_version,
            deployment_name=deployment_name,
            openai_api_type="azure",
            temperature=0.7  # Default temperature
        )
        st.success("Azure OpenAI API connection successful!")
    except Exception as e:
        st.error(f"Error connecting to Azure OpenAI API: {str(e)}")
else:
    st.warning("Please enter your Azure OpenAI API Key.")

# File uploader to accept both .txt and .docx files
uploaded_files = st.file_uploader("Upload one or more transcript files (TXT or Word)", type=["txt", "docx"], accept_multiple_files=True)

# Temperature slider
temperature = st.slider("Set the temperature for the output (0 = deterministic, 1 = creative)", min_value=0.0, max_value=1.0, value=0.7)

# Initialize session state variables
if 'combined_content' not in st.session_state:
    st.session_state['combined_content'] = ""
if 'final_report' not in st.session_state:
    st.session_state['final_report'] = ""

# Button to start processing
if st.button("Generate Research Article"):
    if not uploaded_files:
        st.error("Please upload at least one transcript file.")
    elif not azure_api_key:
        st.error("Please enter your Azure OpenAI API Key.")
    else:
        # Process files
        st.session_state['combined_content'] = ""
        for i, uploaded_file in enumerate(uploaded_files, 1):
            if uploaded_file.type == "text/plain":
                file_content = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_content = read_docx(uploaded_file)

            st.session_state['combined_content'] += f"--- File {i}: {uploaded_file.name} ---\n"
            st.session_state['combined_content'] += file_content + "\n\n"

        # Ensure combined content is not empty
        if st.session_state['combined_content']:
            try:
                # Define agents and tasks
                planner = Agent(
                    role="Content Planner",
                    goal="Plan content based on transcripts.",
                    backstory="Plan a structured research article.",
                    llm=llm,
                    allow_delegation=False,
                    verbose=True,
                    temperature=temperature
                )

                writer = Agent(
                    role="Content Writer",
                    goal="Write a cohesive article based on the plan.",
                    backstory="Write a polished research article.",
                    llm=llm,
                    allow_delegation=False,
                    verbose=True,
                    temperature=temperature
                )

                editor = Agent(
                    role="Editor",
                    goal="Edit and refine the research article.",
                    backstory="Finalize the research article for publication.",
                    llm=llm,
                    allow_delegation=False,
                    verbose=True,
                    temperature=temperature
                )

                # Define tasks
                plan = Task(description="Plan content for the given transcripts.", agent=planner)
                write = Task(description="Write a research article based on the content plan.", agent=writer)
                edit = Task(description="Edit and finalize the research article.", agent=editor)

                # Crew to manage agents and tasks
                crew = Crew(agents=[planner, writer, editor], tasks=[plan, write, edit], verbose=True)

                # Generate report
                with st.spinner("Generating research article... This may take a few minutes."):
                    result = crew.kickoff()

                # Display the final report
                st.session_state['final_report'] = result
                st.success("Research article generated successfully!")
                st.markdown(result)
            except Exception as e:
                st.error(f"Error: {str(e)}")
                st.error(traceback.format_exc())
        else:
            st.error("No content to process.")

# Buttons to download the combined content and final report
if st.session_state['final_report']:
    doc = Document()
    doc.add_paragraph("Industry Insights Report", style='Heading 1')
    for line in st.session_state['final_report'].split('\n'):
        p = doc.add_paragraph(line.strip())
        p.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Final Report",
        data=buffer.getvalue(),
        file_name="research_article.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
